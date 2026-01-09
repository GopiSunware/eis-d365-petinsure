"""Word (DOCX) Exporter for claim documents."""

import logging
import os
from datetime import datetime, timedelta
from uuid import uuid4

from app.config import DATA_DIR
from app.models.document import ExportFormat, ExportResult, ProcessedBatch

logger = logging.getLogger(__name__)


class WordExporter:
    """Generate Word (DOCX) exports for claim documents."""

    async def export(
        self,
        batch: ProcessedBatch,
        template: str = "claim_letter",
        include_raw_data: bool = False,
    ) -> ExportResult:
        """Generate Word export."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed, using simple text export")
            return await self._export_simple(batch, template, include_raw_data)

        export_id = uuid4()
        filename = f"claim_{batch.batch_id}_{template}.docx"
        output_dir = str(DATA_DIR / "exports")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading('Claim Processing Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Batch info
        doc.add_heading('Batch Information', level=1)
        doc.add_paragraph(f'Batch ID: {batch.batch_id}')
        doc.add_paragraph(f'Created: {batch.created_at.strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph(f'Status: {batch.status.value}')

        # Documents
        doc.add_heading('Documents', level=1)
        for doc_item in batch.documents:
            doc.add_paragraph(doc_item.original_filename, style='List Bullet')

        # Claim Data
        if batch.mapped_claim:
            doc.add_heading('Claim Details', level=1)
            claim = batch.mapped_claim
            if isinstance(claim, dict):
                for key, value in claim.items():
                    if value and not key.startswith('_'):
                        doc.add_paragraph(f'{key}: {value}')

        # Billing
        if batch.billing:
            doc.add_heading('Billing Summary', level=1)
            doc.add_paragraph(f'Claim Amount: ${batch.billing.claim_amount:.2f}')
            doc.add_paragraph(f'Deductible Applied: ${batch.billing.deductible_applied:.2f}')
            doc.add_paragraph(f'Final Payout: ${batch.billing.final_payout:.2f}')

        # AI Decision
        if batch.ai_decision:
            doc.add_heading('AI Analysis', level=1)
            doc.add_paragraph(f'Decision: {batch.ai_decision}')
            doc.add_paragraph(f'Confidence: {(batch.ai_confidence or 0) * 100:.0f}%')
            if batch.ai_reasoning:
                doc.add_paragraph('Reasoning:')
                # Format reasoning with line breaks
                for line in batch.ai_reasoning.split('\\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip(), style='List Bullet')

        doc.save(output_path)

        logger.info(f"Generated Word export: {output_path}")

        return ExportResult(
            export_id=export_id,
            batch_id=batch.batch_id,
            format=ExportFormat.WORD,
            filename=filename,
            local_path=output_path,
            file_size=os.path.getsize(output_path),
            expires_at=datetime.utcnow() + timedelta(hours=24),
        )

    async def _export_simple(
        self,
        batch: ProcessedBatch,
        template: str,
        include_raw_data: bool,
    ) -> ExportResult:
        """Simple text export when python-docx is not available."""
        export_id = uuid4()
        filename = f"claim_{batch.batch_id}_{template}.txt"
        output_dir = str(DATA_DIR / "exports")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)

        content = []
        content.append("CLAIM PROCESSING REPORT")
        content.append("")
        content.append(f"Batch ID: {batch.batch_id}")
        content.append(f"Created: {batch.created_at.strftime('%Y-%m-%d %H:%M')}")
        content.append(f"Status: {batch.status.value}")

        with open(output_path, 'w') as f:
            f.write('\n'.join(content))

        return ExportResult(
            export_id=export_id,
            batch_id=batch.batch_id,
            format=ExportFormat.WORD,
            filename=filename,
            local_path=output_path,
            file_size=os.path.getsize(output_path),
            expires_at=datetime.utcnow() + timedelta(hours=24),
        )
